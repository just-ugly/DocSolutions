from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from flask import Flask, request, jsonify
import os

from dify import upload_file_to_dify

app = Flask(__name__)
UPLOAD_FOLDER = 'backend/files'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def create_docx(data: dict):
    """
    将结构化 JSON 自动转为美观的 Word 文档

    :param data: JSON 数据（包含 sections）
    """
    doc = Document()
    title = data.get("title", "")
    output_path = f"{title}.docx"

    # ===== 全局字体（中文友好）=====
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    # ===== 文档主标题 =====
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(20)

    # ===== 正文内容 =====
    sections = data.get("sections", [])
    for section in sections:
        # 二级标题
        heading_p = doc.add_heading(section.get("heading", ""), level=2)
        heading_p.paragraph_format.space_before = Pt(12)
        heading_p.paragraph_format.space_after = Pt(6)

        # 正文段落
        content_p = doc.add_paragraph()
        for block in section.get("content", []):
            run = content_p.add_run(block.get("text", ""))
            run.bold = block.get("bold", False)
            run.font.size = Pt(12)

        content_p.paragraph_format.first_line_indent = Inches(0.3)
        content_p.paragraph_format.line_spacing = 1.5
        content_p.paragraph_format.space_after = Pt(12)

        # ===== WARNING 块（可选）=====
        if section.get("warning"):
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True

            cell = table.rows[0].cells[0]
            cell_p = cell.paragraphs[0]

            # 背景色
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "FFF3CD")
            tc_pr.append(shd)

            # 边框
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "12")
                border.set(qn("w:color"), "D39E00")
                borders.append(border)
            tc_pr.append(borders)

            # 内容（全部粗体）
            run = cell_p.add_run("⚠ 警告\n" + section["warning"])
            run.bold = True
            run.font.size = Pt(12)

            cell_p.paragraph_format.space_before = Pt(6)
            cell_p.paragraph_format.space_after = Pt(6)

        # ===== ERROR 块（可选）=====
        if section.get("error"):
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True

            cell = table.rows[0].cells[0]
            cell_p = cell.paragraphs[0]

            # 背景色
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F8D7DA")
            tc_pr.append(shd)

            # 边框
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "12")
                border.set(qn("w:color"), "A71D2A")
                borders.append(border)
            tc_pr.append(borders)

            # 内容（全部粗体）
            run = cell_p.add_run("⛔ 禁止\n" + section["error"])
            run.bold = True
            run.font.size = Pt(12)

            cell_p.paragraph_format.space_before = Pt(6)
            cell_p.paragraph_format.space_after = Pt(6)

    # ===== 保存 =====
    doc.save(output_path)
    return output_path


@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    filename = file.filename
    save_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(save_path)

    # Upload to Dify
    upload_file_id = upload_file_to_dify(save_path)
    if not upload_file_id:
        return jsonify({'error': 'Failed to upload file to Dify'}), 500

    return jsonify({'message': 'File uploaded successfully', 'file_path': save_path, 'upload_file_id': upload_file_id}), 200

